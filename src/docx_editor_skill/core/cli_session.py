"""CLI-specific session with deterministic element IDs.

In CLI mode, each invocation opens a file, operates, and saves. Element IDs
must be deterministic (based on document order) so they stay stable across
separate CLI calls as long as the document structure doesn't change.

ID scheme:
    para_001, para_002, ...
    table_001, table_002, ...
    run_001, run_002, ...
    cell_001, cell_002, ...
"""

import logging
from typing import Any, Optional, Dict

from docx import Document
from docx.table import Table

from docx_editor_skill.core.session import Session

logger = logging.getLogger(__name__)

# Counters per prefix for deterministic ID generation
_PREFIX_COUNTERS: Dict[str, int] = {}


class CLISession(Session):
    """Session subclass that generates deterministic, ordered element IDs."""

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        # Per-session prefix counters for deterministic IDs
        self._prefix_counters: Dict[str, int] = {}

    def register_object(
        self,
        obj: Any,
        prefix: str = "obj",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> str:
        """Register a docx object with a deterministic, ordered ID.

        Instead of random UUIDs, IDs are sequential: para_001, para_002, etc.
        """
        count = self._prefix_counters.get(prefix, 0) + 1
        self._prefix_counters[prefix] = count
        obj_id = f"{prefix}_{count:03d}"

        self.object_registry[obj_id] = obj

        # Update reverse cache for context lookup
        if hasattr(obj, "_element"):
            self._element_id_cache[id(obj._element)] = obj_id

        if metadata:
            self.element_metadata[obj_id] = metadata

        logger.debug(f"CLI object registered: {obj_id} (type={type(obj).__name__})")
        return obj_id


def build_registry(session: CLISession) -> None:
    """Scan the document and register all elements with deterministic IDs.

    Walks the document body in order, registering paragraphs, tables,
    runs, and table cells so that IDs are stable across CLI invocations.
    """
    doc = session.document

    # Walk body elements in document order
    body = doc.element.body
    para_list = list(doc.paragraphs)
    table_list = list(doc.tables)

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # Find matching Paragraph object
            for p in para_list:
                if p._element is child:
                    session.register_object(p, "para")
                    # Register runs within paragraph
                    for run in p.runs:
                        session.register_object(run, "run")
                    break

        elif tag == "tbl":
            # Find matching Table object
            for t in table_list:
                if t._element is child:
                    session.register_object(t, "table")
                    # Register cells and their contents
                    _register_table_contents(session, t)
                    break

    logger.info(
        f"Registry built: {len(session.object_registry)} elements "
        f"(counters: {dict(session._prefix_counters)})"
    )


def _register_table_contents(session: CLISession, table: Table) -> None:
    """Register all cells, paragraphs, and runs within a table."""
    for row in table.rows:
        for cell in row.cells:
            session.register_object(cell, "cell")
            for para in cell.paragraphs:
                session.register_object(para, "para")
                for run in para.runs:
                    session.register_object(run, "run")


def open_cli_session(file_path: str) -> CLISession:
    """Open a file and create a CLISession with a fully built registry.

    This is the main entry point for CLI commands. It:
    1. Opens the document
    2. Creates a CLISession
    3. Scans and registers all elements with deterministic IDs
    4. Sets up global state so tool functions work

    Args:
        file_path: Path to the .docx file

    Returns:
        CLISession: Ready-to-use session with deterministic IDs
    """
    import os
    from docx_editor_skill.core.global_state import global_state
    from docx_editor_skill.core.session import session_manager

    abs_path = os.path.abspath(file_path)

    if os.path.exists(abs_path):
        doc = Document(abs_path)
    else:
        # New file — create empty document
        parent = os.path.dirname(abs_path)
        if parent and not os.path.exists(parent):
            raise ValueError(f"Parent directory does not exist: {parent}")
        doc = Document()

    session = CLISession(
        session_id="cli",
        document=doc,
        file_path=abs_path,
    )

    # Build the element registry
    build_registry(session)

    # Register in session_manager and set global state
    session_manager.sessions["cli"] = session
    global_state.active_session_id = "cli"
    global_state.active_file = abs_path

    return session
